import re
import os
from typing import List, Dict, Tuple, Optional
from src.config import (
    CHUNK_SIZE_SMALL, CHUNK_SIZE_LARGE, CHUNK_OVERLAP_SMALL, CHUNK_OVERLAP_LARGE,
    PARENT_CHUNK_SIZE, MIN_CHUNK_LENGTH,
)

_TRADITIONAL_TO_SIMPLIFIED = {
    '法則': '法则', '觀': '观', '論': '论', '導': '导', '業': '业',
    '組織': '组织', '領導': '领导', '決策': '决策', '戰略': '战略',
    '規劃': '规划', '控制': '控制', '溝通': '沟通', '協調': '协调',
    '績效': '绩效', '變革': '变革', '創新': '创新', '競爭': '竞争',
    '營銷': '营销', '運營': '运营', '質量': '质量', '效率': '效率',
    '執行力': '执行力', '協作': '协作', '標準化': '标准化',
    '學': '学', '習': '习', '動': '动', '產': '产', '點': '点',
    '經': '经', '濟': '济', '開': '开', '發': '发', '關': '关',
    '係': '系', '實': '实', '踐': '践', '認': '认', '識': '识',
    '體': '体', '現': '现', '種': '种', '類': '类', '結': '结',
    '構': '构', '過': '过', '程': '程', '環': '环', '境': '境',
    '義': '义', '務': '务', '責': '责', '權': '权', '據': '据',
    '書': '书', '籍': '籍', '說': '说', '話': '话', '語': '语',
    '讀': '读', '寫': '写', '計': '计', '劃': '划', '設': '设',
    '備': '备', '術': '术', '技': '技', '條': '条', '件': '件',
    '規': '规', '則': '则', '範': '范', '圍': '围', '標': '标',
    '準': '准', '確': '确', '應': '应', '當': '当', '從': '从',
    '來': '来', '為': '为', '與': '与', '個': '个', '們': '们',
    '時': '时', '間': '间', '長': '长', '較': '较', '最': '最',
    '進': '进', '行': '行', '對': '对', '於': '于', '將': '将',
    '會': '会', '能': '能', '無': '无', '有': '有', '這': '这',
    '那': '那', '裡': '里', '還': '还', '卻': '却', '讓': '让',
    '給': '给', '被': '被', '把': '把', '已': '已', '曾': '曾',
    '電': '电', '網': '网', '數': '数', '據': '据', '資': '资',
    '源': '源', '總': '总', '結': '结', '評': '评', '價': '价',
    '值': '值', '優': '优', '勢': '势', '劣': '劣', '機': '机',
    '會': '会', '威': '威', '脅': '胁', '戰': '战', '術': '术',
    '層': '层', '級': '级', '職': '职', '務': '务', '階': '阶',
    '段': '段', '歷': '历', '史': '史', '發': '发', '展': '展',
    '變': '变', '遷': '迁', '轉': '转', '型': '型', '態': '态',
    '樣': '样', '式': '式', '種': '种', '區': '区', '別': '别',
    '異': '异', '同': '同', '質': '质', '量': '量', '量': '量',
    '效': '效', '益': '益', '損': '损', '失': '失', '增': '增',
    '減': '减', '擴': '扩', '張': '张', '縮': '缩', '收': '收',
    '支': '支', '費': '费', '預': '预', '算': '算', '財': '财',
    '報': '报', '表': '表', '單': '单', '據': '据', '證': '证',
    '據': '据', '調': '调', '查': '查', '研': '研', '究': '究',
    '分': '分', '析': '析', '測': '测', '試': '试', '驗': '验',
    '證': '证', '估': '估', '計': '计', '算': '算', '統': '统',
    '計': '计', '圖': '图', '表': '表', '冊': '册', '卷': '卷',
    '篇': '篇', '章': '章', '節': '节', '項': '项', '條': '条',
    '款': '款', '目': '目', '錄': '录', '編': '编', '著': '著',
    '譯': '译', '審': '审', '訂': '订', '修': '修', '訂': '订',
    '增': '增', '刪': '删', '補': '补', '充': '充', '附': '附',
    '註': '注', '釋': '释', '參': '参', '考': '考', '引': '引',
    '用': '用', '摘': '摘', '要': '要', '關鍵': '关键', '詞': '词',
    '檢': '检', '索': '索', '鏈': '链', '結': '结', '點': '点',
    '網絡': '网络', '節點': '节点', '邊': '边', '緣': '缘',
    '邊緣': '边缘', '極': '极', '限': '限', '約': '约', '束': '束',
    '條件': '条件', '假設': '假设', '推論': '推论', '結論': '结论',
    '證明': '证明', '定義': '定义', '公理': '公理', '定理': '定理',
    '公式': '公式', '函數': '函数', '變量': '变量', '常數': '常数',
    '參數': '参数', '矩陣': '矩阵', '向量': '向量', '空間': '空间',
    '維度': '维度', '坐標': '坐标', '距離': '距离', '角度': '角度',
    '面積': '面积', '體積': '体积', '長度': '长度', '重量': '重量',
    '速度': '速度', '加速度': '加速度', '力': '力', '壓力': '压力',
    '溫度': '温度', '能量': '能量', '功率': '功率', '頻率': '频率',
    '波長': '波长', '振幅': '振幅', '相位': '相位', '周期': '周期',
}

_DOMAIN_WORDS = [
    '科学管理', '行为科学', '人际关系学派', '权变理论', '系统理论',
    '目标管理', '知识管理', '项目管理', '时间管理', '危机管理',
    '组织行为', '组织结构', '组织文化', '组织变革', '组织设计',
    '领导力', '领导风格', '领导理论', '变革型领导', '交易型领导',
    '竞争优势', '核心竞争力', '差异化战略', '成本领先战略',
    '市场细分', '目标市场', '市场定位', '营销组合',
    '需求层次理论', '双因素理论', '期望理论', '公平理论',
    '精益生产', '六西格玛', '敏捷管理', '迭代开发',
    'SWOT分析', '波特五力', '波士顿矩阵', '平衡计分卡',
    '泰勒制', '法约尔原则', '韦伯官僚制', '霍桑实验',
    '马斯洛需求', '赫茨伯格', '德鲁克管理', '彼得原理',
    '帕累托法则', '二八法则', '长尾理论', '蓝海战略',
    '红海战略', '颠覆式创新', '持续改进', '全面质量管理',
    '供应链管理', '客户关系管理', '企业资源规划',
    '关键绩效指标', 'OKR工作法', 'MBO目标管理',
    '知识工作者', '学习型组织', '扁平化组织', '矩阵式组织',
    '事业部制', '直线职能制', '虚拟组织', '无边界组织',
    '战略联盟', '并购重组', '多元化战略', '国际化战略',
    '企业文化', '狼性文化', '工匠精神', '以人为本',
    '以奋斗者为本', '基业长青', '从优秀到卓越',
    '卓有成效的管理者', '管理的实践', '高效能人士',
    '8020法则', '关键跨越', '精益创业', '最小可行产品',
    '管理思想史', '管理学原理', '公共管理', '行政学',
    '雷恩', '贝德安', '任正非', '华为管理法',
    '科克', '柯林斯', '德鲁克', '彼得斯',
    '产品思维', '用户体验', '商业模式画布',
    '增长黑客', '数据驱动', '数字化转型',
]

_jieba_initialized = False


def _init_jieba():
    global _jieba_initialized
    if _jieba_initialized:
        return
    import jieba
    for word in _DOMAIN_WORDS:
        jieba.add_word(word)
    _jieba_initialized = True


def traditional_to_simplified(text: str) -> str:
    result = []
    i = 0
    while i < len(text):
        matched = False
        for length in range(min(4, len(text) - i), 0, -1):
            substr = text[i:i + length]
            if substr in _TRADITIONAL_TO_SIMPLIFIED:
                result.append(_TRADITIONAL_TO_SIMPLIFIED[substr])
                i += length
                matched = True
                break
        if not matched:
            result.append(text[i])
            i += 1
    return ''.join(result)


class ContentRefiner:
    CHAPTER_PATTERNS = [
        re.compile(r'^第[一二三四五六七八九十百千\d]+[章节讲]'),
        re.compile(r'^第[一二三四五六七八九十百千\d]+部分'),
        re.compile(r'^模块[一二三四五六七八九十\d]+'),
        re.compile(r'^Chapter\s+\d+', re.IGNORECASE),
        re.compile(r'^Part\s+\d+', re.IGNORECASE),
    ]

    SECTION_PATTERNS = [
        re.compile(r'^\d+\.\d+'),
        re.compile(r'^[一二三四五六七八九十]、'),
        re.compile(r'^\(\d+\)'),
        re.compile(r'^[（(][一二三四五六七八九十][）)]'),
    ]

    DEFINITION_PATTERNS = [
        re.compile(r'是指\b'),
        re.compile(r'定义为?\b'),
        re.compile(r'核心是\b'),
        re.compile(r'本质是\b'),
        re.compile(r'所谓\b'),
        re.compile(r'指的是\b'),
        re.compile(r'即\b.{2,15}的\b(?:过程|活动|方式|方法|手段|体系|系统)'),
        re.compile(r'就是\b'),
        re.compile(r'可以理解为\b'),
        re.compile(r'被定义为\b'),
    ]

    FRAMEWORK_PATTERNS = [
        re.compile(r'包括以下(?:几个)?(?:方面|部分|要素|维度|层次|类型|阶段)'),
        re.compile(r'主要(?:有|包括|分为|包含)'),
        re.compile(r'分为\b'),
        re.compile(r'包含\b.{0,5}(?:方面|部分|要素|类型|阶段)'),
        re.compile(r'由\b.{2,8}组成'),
        re.compile(r'由\b.{2,8}构成'),
        re.compile(r'(?:三大|四大|五大|六大|七个|八个|九个|十个)\b'),
        re.compile(r'几[个种类型]'),
    ]

    CONCLUSION_PATTERNS = [
        re.compile(r'因此[，,]?\b'),
        re.compile(r'总之[，,]?\b'),
        re.compile(r'综上(?:所述)?[，,]?\b'),
        re.compile(r'可见[，,]?\b'),
        re.compile(r'由此可见\b'),
        re.compile(r'总的来说[，,]?\b'),
        re.compile(r'一言以蔽之\b'),
        re.compile(r'简而言之\b'),
        re.compile(r'这表明\b'),
        re.compile(r'这意味(?:着)?\b'),
    ]

    MANAGEMENT_TERMS = {
        '管理', '组织', '领导', '决策', '战略', '规划', '控制', '激励',
        '沟通', '协调', '授权', '目标', '绩效', '文化', '变革', '创新',
        '竞争', '营销', '人力资源', '财务', '运营', '供应链', '质量',
        '效率', '效能', '执行力', '团队', '协作', '流程', '标准化',
        '泰勒', '法约尔', '韦伯', '德鲁克', '马斯洛', '波特', '梅奥',
        'SWOT', 'KPI', 'OKR', 'MBO', 'TQM', 'JIT', 'MVP',
        '科学管理', '行为科学', '人际关系', '权变理论', '系统理论',
        '目标管理', '知识管理', '项目管理', '时间管理', '危机管理',
        '组织行为', '组织结构', '组织文化', '组织变革',
        '领导力', '领导风格', '领导理论',
        '竞争优势', '核心竞争力', '差异化', '成本领先',
        '市场细分', '目标市场', '市场定位', '4P', '营销组合',
        '需求层次', '双因素', '期望理论', '公平理论',
        '精益', '六西格玛', '敏捷', '迭代',
    }

    NEGATIVE_PATTERNS = [
        re.compile(r'^比如[，,]?\s'),
        re.compile(r'^例如[，,]?\s'),
        re.compile(r'^举例来说'),
        re.compile(r'^以.{1,8}为例'),
        re.compile(r'^如表\d'),
        re.compile(r'^如图\d'),
        re.compile(r'^见表\d'),
        re.compile(r'^见图\d'),
        re.compile(r'^资料来源[：:]'),
    ]

    def __init__(self, max_chars_per_doc: int = 80000):
        self.max_chars_per_doc = max_chars_per_doc

    def refine(self, text: str, doc_name: str = "") -> str:
        if len(text) <= self.max_chars_per_doc:
            return text

        chapters = self._split_into_chapters(text)

        if not chapters:
            return self._fallback_extract(text)

        total_len = sum(len(c["content"]) for c in chapters)
        if total_len <= self.max_chars_per_doc:
            return text

        refined_chapters = []
        for ch in chapters:
            ch_len = len(ch["content"])
            ch_budget = max(500, int(self.max_chars_per_doc * ch_len / total_len))
            refined = self._refine_chapter(ch["content"], ch_budget, ch["title"])
            if ch["title"]:
                refined = ch["title"] + "\n" + refined
            refined_chapters.append(refined)

        result = "\n\n".join(refined_chapters)
        if len(result) > self.max_chars_per_doc * 1.1:
            ratio = self.max_chars_per_doc / len(result)
            new_budgets = []
            for ch in chapters:
                ch_len = len(ch["content"])
                new_budgets.append(max(300, int(self.max_chars_per_doc * ch_len / total_len * ratio)))
            refined_chapters = []
            for ch, budget in zip(chapters, new_budgets):
                refined = self._refine_chapter(ch["content"], budget, ch["title"])
                if ch["title"]:
                    refined = ch["title"] + "\n" + refined
                refined_chapters.append(refined)
            result = "\n\n".join(refined_chapters)

        return result

    def _split_into_chapters(self, text: str) -> List[Dict]:
        lines = text.split('\n')
        chapters = []
        current_title = ""
        current_lines = []
        current_start = 0

        for i, line in enumerate(lines):
            stripped = line.strip()
            is_chapter = False
            for pat in self.CHAPTER_PATTERNS:
                if pat.match(stripped) and len(stripped) < 80:
                    is_chapter = True
                    break

            if is_chapter:
                if current_lines:
                    chapters.append({
                        "title": current_title,
                        "content": "\n".join(current_lines),
                        "start": current_start,
                    })
                current_title = stripped
                current_lines = []
                current_start = i
            else:
                current_lines.append(line)

        if current_lines:
            chapters.append({
                "title": current_title,
                "content": "\n".join(current_lines),
                "start": current_start,
            })

        if len(chapters) == 1 and not chapters[0]["title"]:
            return self._split_by_sections(text)

        return chapters

    def _split_by_sections(self, text: str) -> List[Dict]:
        lines = text.split('\n')
        sections = []
        current_title = ""
        current_lines = []
        chunk_size = 5000

        for i, line in enumerate(lines):
            stripped = line.strip()
            is_section = False
            for pat in self.SECTION_PATTERNS:
                if pat.match(stripped) and len(stripped) < 80:
                    is_section = True
                    break

            if is_section:
                if current_lines:
                    sections.append({
                        "title": current_title,
                        "content": "\n".join(current_lines),
                        "start": 0,
                    })
                current_title = stripped
                current_lines = []
            else:
                current_lines.append(line)

        if current_lines:
            sections.append({
                "title": current_title,
                "content": "\n".join(current_lines),
                "start": 0,
            })

        if not sections:
            for i in range(0, len(text), chunk_size):
                sections.append({
                    "title": "",
                    "content": text[i:i + chunk_size],
                    "start": i,
                })

        return sections

    def _refine_chapter(self, text: str, budget: int, title: str = "") -> str:
        if len(text) <= budget:
            return text

        paragraphs = self._split_into_paragraphs(text)
        if not paragraphs:
            return text[:budget]

        n = len(paragraphs)
        head_count = max(1, int(n * 0.2))
        tail_count = max(1, int(n * 0.1))

        forced_indices = set(range(head_count)) | set(range(max(0, n - tail_count), n))
        forced_len = sum(len(paragraphs[i]) + 1 for i in forced_indices)
        remaining_budget = max(0, budget - forced_len)

        middle_indices = set(range(n)) - forced_indices
        middle_paragraphs = [(i, paragraphs[i]) for i in sorted(middle_indices)]

        scored_middle = []
        prev_is_heading = False
        for i, para in middle_paragraphs:
            prev_idx = i - 1
            prev_is_heading = self._is_heading(paragraphs[prev_idx]) if prev_idx >= 0 else False
            score = self._score_paragraph(para, prev_is_heading)
            scored_middle.append((i, para, score))

        scored_middle.sort(key=lambda x: x[2], reverse=True)

        selected_middle = set()
        selected_middle_len = 0
        for idx, para, score in scored_middle:
            para_len = len(para) + 1
            if selected_middle_len + para_len <= remaining_budget:
                selected_middle.add(idx)
                selected_middle_len += para_len

        all_selected = forced_indices | selected_middle
        result_parts = [paragraphs[i] for i in sorted(all_selected)]

        return "\n\n".join(result_parts)

    def _split_into_paragraphs(self, text: str) -> List[str]:
        raw = re.split(r'\n{2,}', text)
        paragraphs = []
        for p in raw:
            p = p.strip()
            if len(p) < 10:
                continue
            if len(p) > 1500:
                sub_paras = re.split(r'(?<=[。！？；])', p)
                current = ""
                for sp in sub_paras:
                    if len(current) + len(sp) <= 800:
                        current += sp
                    else:
                        if current:
                            paragraphs.append(current.strip())
                        current = sp
                if current:
                    paragraphs.append(current.strip())
            else:
                paragraphs.append(p)
        return paragraphs

    def _score_paragraph(self, para: str, prev_is_heading: bool = False) -> float:
        score = 0.3

        if prev_is_heading:
            score += 0.30

        for pat in self.DEFINITION_PATTERNS:
            if pat.search(para):
                score += 0.35
                break

        for pat in self.FRAMEWORK_PATTERNS:
            if pat.search(para):
                score += 0.25
                break

        for pat in self.CONCLUSION_PATTERNS:
            if pat.search(para):
                score += 0.20
                break

        term_hits = sum(1 for t in self.MANAGEMENT_TERMS if t in para)
        term_density = term_hits / max(len(para) / 100, 1)
        score += min(0.20, term_density * 0.05)

        if self._is_heading(para):
            score += 0.15

        for pat in self.NEGATIVE_PATTERNS:
            if pat.search(para):
                score -= 0.15
                break

        if len(para) < 30:
            score -= 0.10

        digit_count = len(re.findall(r'\d', para))
        if len(para) > 0 and digit_count / len(para) > 0.3:
            score -= 0.15

        return max(0.0, min(1.0, score))

    def _is_heading(self, para: str) -> bool:
        stripped = para.strip()
        if len(stripped) > 60:
            return False
        for pat in self.CHAPTER_PATTERNS:
            if pat.match(stripped):
                return True
        for pat in self.SECTION_PATTERNS:
            if pat.match(stripped):
                return True
        return False

    def _fallback_extract(self, text: str) -> str:
        if len(text) <= self.max_chars_per_doc:
            return text
        return text[:self.max_chars_per_doc]


class ChineseTextSplitter:
    def __init__(self, chunk_size=CHUNK_SIZE_SMALL, chunk_overlap=CHUNK_OVERLAP_SMALL,
                 parent_chunk_size=PARENT_CHUNK_SIZE, min_chunk_length=MIN_CHUNK_LENGTH,
                 granularity="small"):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.parent_chunk_size = parent_chunk_size
        self.min_chunk_length = min_chunk_length
        self.granularity = granularity
        self.separators = [
            "\n\n",
            "\n",
            "。",
            "；",
            "！",
            "？",
            "……",
            "，",
            " ",
            "",
        ]

    @classmethod
    def create_small_splitter(cls) -> "ChineseTextSplitter":
        return cls(
            chunk_size=CHUNK_SIZE_SMALL,
            chunk_overlap=CHUNK_OVERLAP_SMALL,
            parent_chunk_size=PARENT_CHUNK_SIZE,
            granularity="small",
        )

    @classmethod
    def create_large_splitter(cls) -> "ChineseTextSplitter":
        return cls(
            chunk_size=CHUNK_SIZE_LARGE,
            chunk_overlap=CHUNK_OVERLAP_LARGE,
            parent_chunk_size=PARENT_CHUNK_SIZE,
            granularity="large",
        )

    def _clean_text(self, text: str) -> str:
        text = re.sub(r'[\r\t]+', '', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)
        return text.strip()

    def _split_by_separator(self, text: str, separator: str) -> List[str]:
        if separator == "":
            return list(text)
        parts = text.split(separator)
        result = []
        for i, part in enumerate(parts):
            if i < len(parts) - 1:
                part = part + separator
            if part:
                result.append(part)
        return result

    def _merge_splits(self, splits: List[str]) -> List[str]:
        chunks = []
        current = ""
        for split in splits:
            if len(current) + len(split) <= self.chunk_size:
                current += split
            else:
                if current:
                    chunks.append(current)
                if len(split) > self.chunk_size:
                    sub_chunks = self._force_split_long(split)
                    if current and self.chunk_overlap > 0:
                        overlap_text = current[-self.chunk_overlap:]
                        if sub_chunks:
                            sub_chunks[0] = overlap_text + sub_chunks[0]
                    chunks.extend(sub_chunks[:-1])
                    current = sub_chunks[-1] if sub_chunks else ""
                else:
                    current = split
        if current:
            chunks.append(current)
        return chunks

    def _force_split_long(self, text: str) -> List[str]:
        chunks = []
        for i in range(0, len(text), self.chunk_size - self.chunk_overlap):
            chunk = text[i:i + self.chunk_size]
            if chunk:
                chunks.append(chunk)
        return chunks

    def split_text(self, text: str) -> List[str]:
        text = self._clean_text(text)
        return self._recursive_split(text, self.separators)

    def _recursive_split(self, text: str, separators: List[str]) -> List[str]:
        final_chunks = []
        separator = separators[-1]
        new_separators = []

        for i, s in enumerate(separators):
            if s == "":
                separator = s
                break
            if s in text:
                separator = s
                new_separators = separators[i + 1:]
                break

        splits = self._split_by_separator(text, separator)
        good_splits = []

        for s in splits:
            if len(s) <= self.chunk_size:
                good_splits.append(s)
            else:
                if good_splits:
                    merged = self._merge_splits(good_splits)
                    final_chunks.extend(merged)
                    good_splits = []
                if new_separators:
                    recursive_result = self._recursive_split(s, new_separators)
                    final_chunks.extend(recursive_result)
                else:
                    forced = self._force_split_long(s)
                    final_chunks.extend(forced)

        if good_splits:
            merged = self._merge_splits(good_splits)
            final_chunks.extend(merged)

        return final_chunks

    def create_chunks_with_parents(self, text: str, doc_name: str) -> List[Dict]:
        small_chunks = self.split_text(text)
        result = []

        clean_text = self._clean_text(text)
        parent_step = max(self.parent_chunk_size - self.chunk_overlap, 1)
        parent_chunks = []
        for i in range(0, len(clean_text), parent_step):
            parent_chunks.append(clean_text[i:i + self.parent_chunk_size])

        chunk_positions = self._find_chunk_positions(small_chunks, clean_text)

        chapters = _find_all_chapters(clean_text)

        for i, chunk in enumerate(small_chunks):
            stripped = chunk.strip()
            if len(stripped) < self.min_chunk_length:
                continue
            if _is_low_quality_chunk(stripped):
                continue

            pos = chunk_positions[i]
            chunk_center = pos + max(len(chunk) // 2, 1)
            parent_idx = min(chunk_center // parent_step, len(parent_chunks) - 1)
            parent_content = parent_chunks[parent_idx] if parent_idx < len(parent_chunks) else chunk

            chapter_title = ""
            for ch_start, ch_end, ch_title in chapters:
                if ch_start <= pos < ch_end:
                    chapter_title = ch_title
                    break

            result.append({
                "chunk_id": f"{doc_name}_{self.granularity}_chunk_{i}",
                "content": chunk,
                "parent_content": parent_content,
                "source_doc": doc_name,
                "chunk_index": i,
                "position": pos,
                "chapter": chapter_title,
                "granularity": self.granularity,
            })

        return result

    def _find_chunk_positions(self, chunks: List[str], full_text: str) -> List[int]:
        positions = []
        search_start = 0
        for chunk in chunks:
            if not chunk:
                positions.append(search_start)
                continue
            search_key = chunk[:min(80, len(chunk))]
            pos = full_text.find(search_key, search_start)
            if pos != -1:
                positions.append(pos)
                search_start = pos + max(len(chunk), 1)
            else:
                positions.append(search_start)
        return positions


def _is_low_quality_chunk(text: str) -> bool:
    if not text:
        return True
    chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
    total_chars = len(text.strip())
    if total_chars > 0 and chinese_chars / total_chars < 0.2:
        return True
    digit_ratio = len(re.findall(r'\d', text)) / max(total_chars, 1)
    if digit_ratio > 0.6:
        return True
    unique_chars = len(set(text))
    if total_chars > 20 and unique_chars / total_chars < 0.15:
        return True
    return False


class ChunkQualityFilter:
    DEFINITION_PATTERNS_QF = [
        re.compile(r'是指\b'),
        re.compile(r'定义为?\b'),
        re.compile(r'核心是\b'),
        re.compile(r'本质是\b'),
        re.compile(r'所谓\b'),
        re.compile(r'指的是\b'),
        re.compile(r'就是\b'),
    ]

    FRAMEWORK_PATTERNS_QF = [
        re.compile(r'包括以下(?:几个)?(?:方面|部分|要素|维度|类型|阶段)'),
        re.compile(r'主要(?:有|包括|分为|包含)'),
        re.compile(r'分为\b'),
        re.compile(r'由\b.{2,8}组成'),
        re.compile(r'(?:三大|四大|五大|六大|七个|八个|九个|十个)\b'),
    ]

    CONCLUSION_PATTERNS_QF = [
        re.compile(r'因此[，,]?\b'),
        re.compile(r'总之[，,]?\b'),
        re.compile(r'综上(?:所述)?[，,]?\b'),
        re.compile(r'可见[，,]?\b'),
        re.compile(r'总的来说[，,]?\b'),
    ]

    CASE_PATTERNS = [
        re.compile(r'^比如[，,]?\s'),
        re.compile(r'^例如[，,]?\s'),
        re.compile(r'^举例来说'),
        re.compile(r'^以.{1,8}为例'),
        re.compile(r'^案例[：:]'),
    ]

    MANAGEMENT_TERMS_QF = {
        '管理', '组织', '领导', '决策', '战略', '规划', '控制', '激励',
        '沟通', '协调', '授权', '目标', '绩效', '文化', '变革', '创新',
        '竞争', '营销', '人力资源', '财务', '运营', '供应链', '质量',
        '效率', '效能', '执行力', '团队', '协作', '流程', '标准化',
        '泰勒', '法约尔', '韦伯', '德鲁克', '马斯洛', '波特', '梅奥',
        'SWOT', 'KPI', 'OKR', 'MBO', 'TQM', 'JIT', 'MVP',
        '科学管理', '行为科学', '权变理论', '系统理论',
        '目标管理', '知识管理', '项目管理', '时间管理', '危机管理',
        '组织行为', '组织结构', '组织文化', '组织变革',
        '领导力', '领导风格', '领导理论',
        '竞争优势', '核心竞争力', '差异化', '成本领先',
        '市场细分', '目标市场', '市场定位', '4P', '营销组合',
        '需求层次', '双因素', '期望理论', '公平理论',
        '精益', '六西格玛', '敏捷', '迭代',
    }

    def __init__(self, min_term_density: float = 0.005, jaccard_threshold: float = 0.75):
        self.min_term_density = min_term_density
        self.jaccard_threshold = jaccard_threshold

    def _has_theory_content(self, text: str) -> bool:
        for pat in self.DEFINITION_PATTERNS_QF:
            if pat.search(text):
                return True
        for pat in self.FRAMEWORK_PATTERNS_QF:
            if pat.search(text):
                return True
        for pat in self.CONCLUSION_PATTERNS_QF:
            if pat.search(text):
                return True
        return False

    def _is_pure_case(self, text: str) -> bool:
        stripped = text.strip()
        for pat in self.CASE_PATTERNS:
            if pat.match(stripped):
                return not self._has_theory_content(text)
        return False

    def _compute_term_density(self, text: str) -> float:
        if not text:
            return 0.0
        hits = sum(1 for t in self.MANAGEMENT_TERMS_QF if t in text)
        return hits / max(len(text) / 100, 1)

    def _count_sentences(self, text: str) -> int:
        return len(re.findall(r'[。！？；]', text))

    def _jaccard(self, text_a: str, text_b: str) -> float:
        import jieba
        tokens_a = set(w for w in jieba.cut(text_a) if len(w.strip()) >= 2)
        tokens_b = set(w for w in jieba.cut(text_b) if len(w.strip()) >= 2)
        if not tokens_a or not tokens_b:
            return 0.0
        return len(tokens_a & tokens_b) / len(tokens_a | tokens_b)

    def filter_chunks(self, chunks: List[Dict]) -> List[Dict]:
        filtered = []
        for chunk in chunks:
            content = chunk.get("content", "")
            is_case = self._is_pure_case(content)
            if is_case:
                chunk["quality_weight"] = 0.7
                filtered.append(chunk)
                continue
            if not self._has_theory_content(content):
                term_density = self._compute_term_density(content)
                if term_density < self.min_term_density:
                    continue
            if self._count_sentences(content) < 2:
                chinese_ratio = len(re.findall(r'[\u4e00-\u9fff]', content)) / max(len(content), 1)
                if chinese_ratio < 0.3:
                    continue
            chunk["quality_weight"] = 1.0
            filtered.append(chunk)

        return filtered


NOISE_PATTERNS = [
    re.compile(r'ISBN[=\s]*[\d\-Xx]+', re.IGNORECASE),
    re.compile(r'书\s*号[：:]\s*ISBN', re.IGNORECASE),
    re.compile(r'SS号[=＝]\d+', re.IGNORECASE),
    re.compile(r'出版日期[=＝]\d{4}', re.IGNORECASE),
    re.compile(r'出版社[=＝]', re.IGNORECASE),
    re.compile(r'定价[：:]\s*[\d.]+元', re.IGNORECASE),
    re.compile(r'字数[=＝]\d+', re.IGNORECASE),
    re.compile(r'\[General\s+Information\]', re.IGNORECASE),
    re.compile(r'版权所有[,，]\s*翻[印版]必究'),
    re.compile(r'^未经.{0,10}允许[,，].{0,10}不得.{0,10}复制'),
    re.compile(r'Table\s+of\s+Contents', re.IGNORECASE),
    re.compile(r'^图书在版编目', re.IGNORECASE),
    re.compile(r'^CIP\s*数据', re.IGNORECASE),
    re.compile(r'责任编辑[：:]', re.IGNORECASE),
    re.compile(r'^见习编辑[：:]', re.IGNORECASE),
    re.compile(r'^装帧设计[：:]', re.IGNORECASE),
    re.compile(r'^出版发行[：:]', re.IGNORECASE),
    re.compile(r'^客服热线[：:]', re.IGNORECASE),
    re.compile(r'^客服信箱[：:]', re.IGNORECASE),
    re.compile(r'^官方网址[：:]', re.IGNORECASE),
    re.compile(r'^新浪微博\s*@', re.IGNORECASE),
    re.compile(r'^微信公众号[：:]', re.IGNORECASE),
    re.compile(r'^微信号[：:]', re.IGNORECASE),
    re.compile(r'^Http://', re.IGNORECASE),
    re.compile(r'^邮编[：:]\s*\d{6}', re.IGNORECASE),
    re.compile(r'^地址[：:]\s*[\u4e00-\u9fff]+\d+号', re.IGNORECASE),
    re.compile(r'本书由.{0,6}"[^"]*".{0,20}整理', re.IGNORECASE),
    re.compile(r'加小编微信或QQ[：:]', re.IGNORECASE),
    re.compile(r'小编也和结交一些喜欢读书', re.IGNORECASE),
    re.compile(r'关注小编个人微信公众号', re.IGNORECASE),
    re.compile(r'电子书下载网站', re.IGNORECASE),
    re.compile(r'周读\s+网址[：:]', re.IGNORECASE),
    re.compile(r'ireadweek\.com', re.IGNORECASE),
    re.compile(r'^版本图书馆', re.IGNORECASE),
    re.compile(r'^责任校对[：:]', re.IGNORECASE),
    re.compile(r'^封面设计[：:]', re.IGNORECASE),
    re.compile(r'^版次[：:]\s*\d{4}', re.IGNORECASE),
    re.compile(r'^印次[：:]\s*\d{4}', re.IGNORECASE),
    re.compile(r'^印张[：:]', re.IGNORECASE),
    re.compile(r'^开本[：:]', re.IGNORECASE),
    re.compile(r'^书\s*名[：:]', re.IGNORECASE),
    re.compile(r'^作\s*者[：:]', re.IGNORECASE),
    re.compile(r'^译\s*者[：:]', re.IGNORECASE),
    re.compile(r'^编\s*著[：:]', re.IGNORECASE),
    re.compile(r'^著\s*者[：:]', re.IGNORECASE),
    re.compile(r'^I\s*S\s*B\s*N', re.IGNORECASE),
    re.compile(r'^侵权必究', re.IGNORECASE),
    re.compile(r'^华章分社', re.IGNORECASE),
    re.compile(r'^华章数媒', re.IGNORECASE),
    re.compile(r'^华章电子书', re.IGNORECASE),
    re.compile(r'^全球范围内制作与发行', re.IGNORECASE),
    re.compile(r'^北京华章图文', re.IGNORECASE),
    re.compile(r'^北京奥维博世', re.IGNORECASE),
    re.compile(r'^机械工业出版社', re.IGNORECASE),
    re.compile(r'^古吴轩出版社', re.IGNORECASE),
    re.compile(r'^ISBN\s*978', re.IGNORECASE),
    re.compile(r'^\d{4}年\d{1,2}月.{0,5}版$', re.IGNORECASE),
    re.compile(r'^\d{4}年\d{1,2}月.{0,5}次印刷$', re.IGNORECASE),
]

CHAPTER_PATTERNS = [
    re.compile(r'^第[一二三四五六七八九十百千\d]+[章节讲]'),
    re.compile(r'^第[一二三四五六七八九十百千\d]+部分'),
    re.compile(r'^模块[一二三四五六七八九十\d]+'),
    re.compile(r'^\d{1,2}\s'),
    re.compile(r'^Chapter\s+\d+', re.IGNORECASE),
    re.compile(r'^Part\s+\d+', re.IGNORECASE),
    re.compile(r'^导读'),
    re.compile(r'^前言'),
    re.compile(r'^序言'),
    re.compile(r'^绪论'),
]

REPETITIVE_HEADERS = [
    '中国人民大学出版社',
    '高等教育出版社',
    '清华大学出版社',
    '中信出版社',
    '北京大学出版社',
]


def preprocess_text(text: str) -> str:
    text = traditional_to_simplified(text)
    lines = text.split('\n')

    cleaned_lines = []
    prev_line = None

    for i, line in enumerate(lines):
        stripped = line.strip()

        if len(stripped) == 0:
            if cleaned_lines and cleaned_lines[-1] != '':
                cleaned_lines.append('')
            elif cleaned_lines and cleaned_lines[-1] == '':
                pass
            else:
                cleaned_lines.append('')
            continue

        if _is_noise_line(stripped):
            continue

        if len(stripped) <= 1:
            continue
        if len(stripped) == 2 and not re.match(r'^[\u4e00-\u9fff]{2}$', stripped):
            continue

        if stripped == prev_line:
            continue

        if _is_standalone_number(stripped, lines, i):
            continue

        cleaned_lines.append(stripped)
        prev_line = stripped

    text = '\n'.join(cleaned_lines)

    text = re.sub(r'\n{4,}', '\n\n\n', text)
    text = re.sub(r' {2,}', ' ', text)

    return text.strip()


def _is_noise_line(line: str) -> bool:
    for pattern in NOISE_PATTERNS:
        if pattern.search(line):
            return True
    if line.rstrip() in REPETITIVE_HEADERS:
        return True
    return False


def _is_standalone_number(line: str, all_lines: list, idx: int) -> bool:
    stripped = line.strip()
    if not re.match(r'^\d{1,4}$', stripped):
        return False
    prev_empty = idx == 0 or not all_lines[idx - 1].strip()
    next_empty = idx >= len(all_lines) - 1 or not all_lines[idx + 1].strip()
    return prev_empty and next_empty


def extract_content(text: str, max_chars: int) -> str:
    if len(text) <= max_chars:
        return text

    clean = preprocess_text(text)

    if len(clean) <= max_chars:
        return clean

    chapter_start = _find_first_chapter(clean)
    body_text = clean[chapter_start:]

    if len(body_text) <= max_chars:
        return body_text

    return _truncate_at_boundary(body_text, max_chars)


def _find_first_chapter(text: str) -> int:
    lines = text.split('\n')
    chapter_candidates = []
    pos = 0
    for i, line in enumerate(lines):
        stripped = line.strip()
        for pattern in CHAPTER_PATTERNS:
            if pattern.match(stripped):
                chapter_candidates.append((pos, i, stripped))
                break
        pos += len(line) + 1

    for ch_pos, ch_idx, ch_text in chapter_candidates:
        has_body = False
        for j in range(ch_idx + 1, min(ch_idx + 8, len(lines))):
            ns = lines[j].strip()
            if len(ns) > 50 and not _is_noise_line(ns):
                has_body = True
                break
            if len(ns) > 30:
                is_ch = any(p.match(ns) for p in CHAPTER_PATTERNS)
                if not is_ch:
                    has_body = True
                    break
        if has_body:
            return ch_pos

    pos = 0
    for i, line in enumerate(lines):
        stripped = line.strip()
        if len(stripped) > 20 and not _is_noise_line(stripped):
            return pos
        pos += len(line) + 1
    return 0


def _find_all_chapters(text: str) -> List[Tuple[int, int, str]]:
    lines = text.split('\n')
    chapter_starts = []
    pos = 0
    for i, line in enumerate(lines):
        stripped = line.strip()
        for pattern in CHAPTER_PATTERNS:
            if pattern.match(stripped) and len(stripped) < 80:
                chapter_starts.append((pos, stripped[:60]))
                break
        pos += len(line) + 1

    if not chapter_starts:
        return []

    chapters = []
    for i, (start, title) in enumerate(chapter_starts):
        end = chapter_starts[i + 1][0] if i + 1 < len(chapter_starts) else len(text)
        chapters.append((start, end, title))

    return chapters


def _truncate_at_boundary(text: str, max_chars: int) -> str:
    if len(text) <= max_chars:
        return text
    head = text[:max_chars]
    cut_points = [head.rfind(s) for s in ['\n\n', '\n', '。', '；', '！', '？']]
    valid = [c for c in cut_points if c >= 0]
    if not valid:
        return head
    cut = max(valid)
    if cut > max_chars * 0.6:
        return text[:cut + 1]
    return head


def _clean_doc_name(filename: str) -> str:
    name = os.path.splitext(filename)[0]
    patterns = [
        r'\s*[-_]\s*z-library[^\s]*',
        r'\s*[-_]\s*z-lib[^\s]*',
        r'\s*[-_]\s*1lib[^\s]*',
        r'\s*[-_]\s*z\.library[^\s]*',
        r'\s*\(\s*z-library[^)]*\)',
        r'\s*\(\s*z-lib[^)]*\)',
        r'\s*\(\s*1lib[^)]*\)',
        r'\s*\(\s*etc\.[^)]*\)',
        r'\s*[-_]\s*etc[^_]*',
    ]
    for p in patterns:
        name = re.sub(p, '', name, flags=re.IGNORECASE)
    name = name.strip()
    name = re.sub(r'[-_]{2,}', '-', name)
    for sep in ['（', '(', '【', '[']:
        idx = name.find(sep)
        if idx > 2:
            name = name[:idx]
    name = name.strip('-_ ')
    name = traditional_to_simplified(name)
    return name if name else filename


def _read_docx_via_zip_xml(filepath: str) -> str:
    """从docx的zip包中直接读取word/document.xml，处理扫描版PDF转DOCX的情况"""
    import zipfile
    from xml.etree import ElementTree as ET

    with zipfile.ZipFile(filepath, 'r') as zf:
        if 'word/document.xml' not in zf.namelist():
            return ""
        xml_content = zf.read('word/document.xml')

    try:
        root = ET.fromstring(xml_content)
    except ET.ParseError:
        return ""

    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    paragraphs = []

    for p in root.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
        texts = []
        for t in p.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
            if t.text:
                texts.append(t.text)
        if texts:
            paragraphs.append(''.join(texts))

    return '\n\n'.join(paragraphs)


def _is_scanned_docx(filepath: str) -> bool:
    """检测DOCX是否为扫描版PDF转制（包含大量图片但无文本）"""
    import zipfile
    try:
        with zipfile.ZipFile(filepath, 'r') as zf:
            files = zf.namelist()
            image_count = len([f for f in files if f.startswith('word/media/image')])
            return image_count > 50  # 超过50张图片认为是扫描版
    except Exception:
        return False


def _ocr_docx_images(filepath: str) -> str:
    """对扫描版DOCX中的图片进行OCR提取文本"""
    import zipfile
    import tempfile
    import shutil
    try:
        from PIL import Image
        import pytesseract
    except ImportError:
        return ""

    text_parts = []
    with zipfile.ZipFile(filepath, 'r') as zf:
        image_files = sorted([f for f in zf.namelist() if f.startswith('word/media/image') and f.endswith(('.jpeg', '.jpg', '.png'))])

        with tempfile.TemporaryDirectory() as tmpdir:
            for img_path in image_files:
                try:
                    zf.extract(img_path, tmpdir)
                    full_path = os.path.join(tmpdir, img_path)
                    img = Image.open(full_path)
                    # OCR识别中文
                    text = pytesseract.image_to_string(img, lang='chi_sim+eng')
                    if text.strip():
                        text_parts.append(text.strip())
                except Exception:
                    continue

    return "\n\n".join(text_parts)


def _read_docx(filepath: str) -> str:
    errors = []

    # 预检：是否为扫描版DOCX
    is_scanned = _is_scanned_docx(filepath)
    if is_scanned:
        print(f"    [WARN] 检测到扫描版DOCX，尝试OCR提取...")
        try:
            ocr_text = _ocr_docx_images(filepath)
            if len(ocr_text.strip()) > 500:
                print(f"    [OCR] 成功提取 {len(ocr_text)} 字符")
                return ocr_text
            errors.append(f"OCR提取文本过少({len(ocr_text.strip())}字符)，可能需要安装tesseract")
        except Exception as e:
            errors.append(f"OCR失败: {e}")

    # 尝试1: python-docx（标准方式）
    try:
        from docx import Document
        doc = Document(filepath)
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        for table in doc.tables:
            table_lines = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                line = " | ".join(c for c in cells if c)
                if line:
                    table_lines.append(line)
            if table_lines:
                paragraphs.append("\n".join(table_lines))
        result = "\n\n".join(paragraphs)
        if len(result.strip()) > 100:
            return result
        errors.append(f"python-docx提取文本过少({len(result.strip())}字符)")
    except Exception as e:
        errors.append(f"python-docx: {e}")

    # 尝试2: docx2txt
    try:
        import docx2txt
        result = docx2txt.process(filepath)
        if result and len(result.strip()) > 100:
            return result
        errors.append(f"docx2txt提取文本过少({len(result.strip()) if result else 0}字符)")
    except Exception as e:
        errors.append(f"docx2txt: {e}")

    # 尝试3: 直接读取zip中的XML（处理扫描版PDF转DOCX）
    try:
        result = _read_docx_via_zip_xml(filepath)
        if len(result.strip()) > 100:
            return result
        errors.append(f"zip-xml提取文本过少({len(result.strip())}字符)")
    except Exception as e:
        errors.append(f"zip-xml: {e}")

    raise RuntimeError(f"无法解析DOCX {os.path.basename(filepath)}: {'; '.join(errors)}")


def _read_pdf(filepath: str) -> str:
    text = ""

    # 尝试1: PyPDF2
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(filepath)
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
    except Exception:
        pass

    # 尝试2: pdfplumber（对扫描版PDF更好）
    if len(text.strip()) < 100:
        try:
            import pdfplumber
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        text += t + "\n"
        except Exception:
            pass

    # 尝试3: pymupdf (fitz) - 对复杂PDF支持最好
    if len(text.strip()) < 100:
        try:
            import fitz
            doc = fitz.open(filepath)
            for page in doc:
                text += page.get_text() + "\n"
            doc.close()
        except Exception:
            pass

    return text.strip()


def _read_text(filepath: str) -> str:
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()


PARSERS = {
    '.docx': _read_docx,
    '.pdf': _read_pdf,
    '.txt': _read_text,
    '.md': _read_text,
}


def _content_similarity(text_a: str, text_b: str, sample_len: int = 2000) -> float:
    sample_a = text_a[:sample_len]
    sample_b = text_b[:sample_len]
    if not sample_a or not sample_b:
        return 0.0

    def _char_ngrams(text, n=3):
        return set(text[i:i+n] for i in range(len(text) - n + 1))

    ngrams_a = _char_ngrams(sample_a)
    ngrams_b = _char_ngrams(sample_b)
    if not ngrams_a or not ngrams_b:
        return 0.0

    intersection = len(ngrams_a & ngrams_b)
    union = len(ngrams_a | ngrams_b)
    jaccard = intersection / union if union > 0 else 0.0

    containment_a = intersection / len(ngrams_a) if ngrams_a else 0
    containment_b = intersection / len(ngrams_b) if ngrams_b else 0

    if containment_a > 0.80 or containment_b > 0.80:
        return 0.85

    return jaccard


def load_knowledge_base(kb_dir: str, use_refiner: bool = True) -> List[Dict]:
    from src.config import REFINER_MAX_CHARS_PER_DOC

    _init_jieba()

    documents = []
    skipped = []
    failed = []

    all_files = sorted(os.listdir(kb_dir))
    file_dict: Dict[str, List[Tuple[str, str, str]]] = {}
    for filename in all_files:
        _, ext = os.path.splitext(filename)
        ext = ext.lower()
        filepath = os.path.join(kb_dir, filename)

        if ext == '.doc' and not filename.lower().endswith('.docx'):
            skipped.append(f"{filename} (旧版.doc格式，需antiword)")
            continue

        if ext not in PARSERS:
            skipped.append(f"{filename} (不支持格式: {ext})")
            continue

        clean_name = _clean_doc_name(filename)
        if clean_name not in file_dict:
            file_dict[clean_name] = []
        file_dict[clean_name].append((filename, ext, filepath))

    FORMAT_PRIORITY = {'.md': 0, '.txt': 1, '.docx': 2, '.pdf': 3}
    deduped_files = []
    for clean_name, variants in file_dict.items():
        if len(variants) > 1:
            variants.sort(key=lambda x: FORMAT_PRIORITY.get(x[1], 99))
            chosen = variants[0]
            dup_names = [v[0][:40] for v in variants[1:]]
            print(f"[DEDUP] {clean_name}: 保留 {chosen[1]} 版本，跳过重复 {dup_names}")
            deduped_files.append(chosen)
        else:
            deduped_files.append(variants[0])

    for filename, ext, filepath in deduped_files:
        try:
            content = PARSERS[ext](filepath)
            content = preprocess_text(content)
            if len(content) < 500:
                skipped.append(f"{filename} (预处理后文本量<500字符)")
                continue
            documents.append({
                "doc_name": _clean_doc_name(filename),
                "content": content,
            })
        except Exception as e:
            failed.append(f"{filename} ({e})")

    content_deduped = []
    removed_by_content = []
    for i, doc in enumerate(documents):
        is_dup = False
        for j, existing in enumerate(content_deduped):
            sim = _content_similarity(doc["content"], existing["content"])
            if sim > 0.70:
                is_dup = True
                removed_by_content.append(
                    f'{doc["doc_name"]} (与 {existing["doc_name"]} 相似度={sim:.2f})'
                )
                break
        if not is_dup:
            content_deduped.append(doc)

    if removed_by_content:
        print(f"[DEDUP-CONTENT] 基于内容相似度去除 {len(removed_by_content)} 个重复文档:")
        for r in removed_by_content:
            print(f"  - {r}")

    if skipped:
        print(f"[SKIP] 跳过 {len(skipped)} 个文件: {[s[:40] for s in skipped]}")
    if failed:
        print(f"[FAIL] 读取失败 {len(failed)} 个文件: {failed}")

    if use_refiner:
        refiner = ContentRefiner(max_chars_per_doc=REFINER_MAX_CHARS_PER_DOC)
        total_before = sum(len(d["content"]) for d in content_deduped)
        for doc in content_deduped:
            original_len = len(doc["content"])
            doc["content"] = refiner.refine(doc["content"], doc["doc_name"])
            refined_len = len(doc["content"])
            if original_len != refined_len:
                print(f"  [REFINE] {doc['doc_name'][:30]}: {original_len//1000}K -> {refined_len//1000}K ({refined_len*100//original_len}%)")
        total_after = sum(len(d["content"]) for d in content_deduped)
        print(f"[REFINE] 总计: {total_before//1000}K -> {total_after//1000}K ({total_after*100//total_before}%)")

    return content_deduped


def select_dense_chunks(all_chunks: List[Dict], max_dense_chunks: int = 0) -> List[Dict]:
    if max_dense_chunks <= 0 or len(all_chunks) <= max_dense_chunks:
        return all_chunks

    doc_chunks: Dict[str, List[Dict]] = {}
    for chunk in all_chunks:
        doc = chunk["source_doc"]
        if doc not in doc_chunks:
            doc_chunks[doc] = []
        doc_chunks[doc].append(chunk)

    total = len(all_chunks)
    selected = []

    for doc, chunks in doc_chunks.items():
        budget = max(1, int(max_dense_chunks * len(chunks) / total))
        if len(chunks) <= budget:
            selected.extend(chunks)
            continue

        chapter_groups: Dict[str, List[Dict]] = {}
        no_chapter = []
        for c in chunks:
            ch = c.get("chapter", "")
            if ch:
                if ch not in chapter_groups:
                    chapter_groups[ch] = []
                chapter_groups[ch].append(c)
            else:
                no_chapter.append(c)

        all_groups = list(chapter_groups.values())
        if no_chapter:
            all_groups.append(no_chapter)

        num_groups = len(all_groups)
        base_per_group = max(1, budget // num_groups)
        remaining = budget - base_per_group * num_groups

        doc_selected = []
        for gi, group in enumerate(all_groups):
            group_budget = base_per_group + (1 if gi < remaining else 0)
            if len(group) <= group_budget:
                doc_selected.extend(group)
            else:
                step = len(group) / group_budget
                for j in range(group_budget):
                    idx = int(j * step)
                    doc_selected.append(group[idx])

        selected.extend(doc_selected)

    return selected
